#!/usr/bin/env python
"""Post-process an MDPI transplant DOCX so the real front matter lands in the
template's MDPI_1.x styles instead of being duplicated as MDPI_1.8_keywords body.

The html2doc transplant keeps the template's placeholder front matter (Title,
Firstname Lastname, Abstract) and dumps the HTML's own title/author/abstract
block as keyword-styled paragraphs before the Introduction. This script reads
the authoritative front matter from the source HTML (by CSS class) and:
  - fills MDPI_1.1_article_type / 1.2_title / 1.3_authornames / 1.6_affiliation
    / 1.7_abstract placeholders with the real text,
  - inserts a real MDPI_1.8_keywords line after the abstract (if the paper has
    keywords),
  - deletes the duplicated keyword-styled run and the editorial "Draft for
    MDPI..." note.

Usage: python _mdpi_frontmatter.py <source.html> <built.docx>
"""
import sys, pathlib
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn

def parse_front(html_path):
    soup = BeautifulSoup(pathlib.Path(html_path).read_text(encoding="utf-8", errors="ignore"),
                         "html.parser")
    body = soup.body or soup
    fm = {"art": None, "title": None, "authors": None, "affils": [], "corr": None,
          "abstract": None, "keywords": None}
    for el in body.find_all(["div", "h1", "p"]):
        cls = set(el.get("class") or [])
        txt = " ".join(el.get_text(" ", strip=True).split())   # collapse all whitespace
        if not txt:
            continue
        if "artType" in cls and fm["art"] is None:
            fm["art"] = txt
        elif el.name == "h1" and fm["title"] is None:
            fm["title"] = txt
        elif "authors" in cls and fm["authors"] is None:
            fm["authors"] = txt
        elif "affil" in cls:
            fm["affils"].append(txt)
        elif "corr" in cls and fm["corr"] is None:
            fm["corr"] = txt
        elif "abstract" in cls and fm["abstract"] is None:
            fm["abstract"] = txt
        elif "keywords" in cls and fm["keywords"] is None:
            fm["keywords"] = txt
        if el.name == "h2":
            break
    # strip leading labels; they are re-added bold
    for k in ("abstract", "keywords"):
        if fm[k]:
            lab = k.capitalize() + ":"
            if fm[k].lower().startswith(lab.lower()):
                fm[k] = fm[k][len(lab):].strip()
    return fm

def set_runs(p, segments):
    """Replace a paragraph's runs with (text, bold) segments, keeping its style."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold

def first_by_style(doc, style_name):
    for p in doc.paragraphs:
        if p.style.name == style_name:
            return p
    return None

def all_by_style(doc, style_name):
    return [p for p in doc.paragraphs if p.style.name == style_name]

def main():
    html_path, docx_path = sys.argv[1], sys.argv[2]
    fm = parse_front(html_path)
    doc = Document(docx_path)

    # locate body start (first real section heading)
    body_start_el = None
    for p in doc.paragraphs:
        if p.style.name == "MDPI_2.1_heading1":
            body_start_el = p._element
            break

    # 1) fill scalar placeholders
    if fm["art"]:
        p = first_by_style(doc, "MDPI_1.1_article_type")
        if p: set_runs(p, [(fm["art"], False)])
    if fm["title"]:
        p = first_by_style(doc, "MDPI_1.2_title")
        if p: set_runs(p, [(fm["title"], False)])
    if fm["authors"]:
        p = first_by_style(doc, "MDPI_1.3_authornames")
        if p: set_runs(p, [(fm["authors"], False)])

    # 2) affiliations + correspondence -> the MDPI_1.6_affiliation placeholders
    aff_ps = all_by_style(doc, "MDPI_1.6_affiliation")
    aff_lines = list(fm["affils"]) + ([fm["corr"]] if fm["corr"] else [])
    for i, ptext in enumerate(aff_lines):
        if i < len(aff_ps):
            set_runs(aff_ps[i], [(ptext, False)])
    for extra in aff_ps[len(aff_lines):]:          # remove unused placeholders
        extra._element.getparent().remove(extra._element)

    # 3) abstract placeholder
    abs_p = first_by_style(doc, "MDPI_1.7_abstract")
    if abs_p is not None:
        if fm["abstract"]:
            set_runs(abs_p, [("Abstract: ", True), (fm["abstract"], False)])
        else:
            abs_p._element.getparent().remove(abs_p._element)  # SI: no abstract
            abs_p = None

    # 4) delete the duplicated keyword-styled donor run (everything before body
    #    start that is MDPI_1.8_keywords)
    for p in list(doc.paragraphs):
        if body_start_el is not None and p._element is body_start_el:
            break
        if p.style.name == "MDPI_1.8_keywords":
            p._element.getparent().remove(p._element)

    # 5) insert the real keywords line right after the abstract (main paper only)
    if fm["keywords"]:
        anchor = abs_p if abs_p is not None else first_by_style(doc, "MDPI_1.3_authornames")
        kw_p = doc.add_paragraph(style="MDPI_1.8_keywords")
        set_runs(kw_p, [("Keywords: ", True), (fm["keywords"], False)])
        if anchor is not None:
            anchor._element.addnext(kw_p._element)

    doc.save(docx_path)
    print(f"front matter fixed: {docx_path}")
    print(f"  art='{fm['art']}' title~'{(fm['title'] or '')[:40]}...' "
          f"affils={len(fm['affils'])} abstract={'yes' if fm['abstract'] else 'no'} "
          f"keywords={'yes' if fm['keywords'] else 'no'}")

if __name__ == "__main__":
    main()
